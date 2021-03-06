"""
application.py
- creates a Flask app instance and registers the database object.
"""

import enum
from datetime import datetime

from docx import Document
from docx.shared import Mm
from flask import Flask, Blueprint, jsonify, g, request, send_file
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from flask_httpauth import HTTPBasicAuth
from passlib.apps import custom_app_context as pwd_context
from itsdangerous import (TimedJSONWebSignatureSerializer as Serializer, BadSignature, SignatureExpired)

app = Flask(__name__)
CORS(app, expose_headers=["x-suggested-filename"])
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///../adc.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["SECRET_KEY"] = "desidragons"
db = SQLAlchemy(app)

api = Blueprint('api', __name__)

auth = HTTPBasicAuth()

# ---------------- Association tables -------------------- #

student_module_table = db.Table("student_to_module",
        db.Column("student_id", db.Integer, db.ForeignKey("student.id"), primary_key=True),
        db.Column("module_id", db.Integer, db.ForeignKey("module.id"), primary_key=True))

# ---------------- Models -------------------- #

class User(db.Model):  
    __tablename__ = "user"

    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    name = db.Column(db.String(80), nullable=False)
    usertype = db.Column(db.String(80), nullable=False)
    email = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False, default="$6$rounds=656000$Mekxk00d5L5we4/3$CwY2NGNDHX1Yt6khinGYUJm/I5s2.3bmjgMLmmsRIVzowR.LWPqIF2KrauyfjuzyWD2MrPToRVNHxYCZAPsFf1")

    def to_dict(self):
        return dict(id=self.id,
                    name=self.name,
                    usertype=str(self.usertype),
                    email=self.email)

    def hash_password(self, password):
        self.password_hash = pwd_context.encrypt(password)

    def verify_password(self, password):
        return pwd_context.verify(password, self.password_hash)

    def generate_auth_token(self, expiration=600):
        s = Serializer(app.config['SECRET_KEY'], expires_in=expiration)
        return s.dumps({ 'id': self.id })

    @staticmethod
    def verify_auth_token(token):
        s = Serializer(app.config['SECRET_KEY'])
        try:
            data = s.loads(token)
        except SignatureExpired:
            return None # valid token, but expired.
        except BadSignature:
            return None # invalid token.
        user = User.query.get(data['id'])
        return user

class Module(db.Model):  
    __tablename__ = "module"

    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    code = db.Column(db.String(10), nullable=False)
    name = db.Column(db.String(80), nullable=False)
    semester = db.Column(db.String(80), nullable=False)
    level = db.Column(db.String(10), nullable=False)
    academic = db.Column(db.Integer, db.ForeignKey('user.id')) # foreign key.
    students = db.relationship(
            "Student",
            secondary=student_module_table,
            back_populates="modules") # many to many.

    def to_dict(self):
        academic_name = User.query.filter_by(id=self.academic).first().name
        return dict(id=self.id,
                    code=self.code,
                    name=self.name,
                    semester=str(self.semester),
                    level=self.level,
                    academic_name=academic_name,
                    academic_id=self.academic,
                    students=[s.id for s in self.students])

class Assessment(db.Model):
    __tablename__ = "assessment"

    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    format = db.Column(db.String(80), nullable=False)
    name = db.Column(db.String(80), nullable=False)
    marks = db.Column(db.Float, nullable=False)
    release_date = db.Column(db.DateTime, nullable=False)
    submission_date = db.Column(db.DateTime, nullable=False)
    form_id = db.Column(db.Integer, db.ForeignKey('form.id'), nullable=False) # foreign key.

    def to_dict(self):
        return dict(id=self.id,
                    asm_format=str(self.format),
                    asm_name=self.name,
                    asm_per=self.marks,
                    asm_release=self.release_date.strftime("%d/%m/%Y"),
                    asm_due=self.submission_date.strftime("%d/%m/%Y"))

class Student(db.Model):
    __tablename__ = "student"

    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    modules = db.relationship(
            "Module",
            secondary=student_module_table,
            back_populates="students") # many to many.

    def to_dict(self):
        return dict(id=self.id,
                    modules=[m.id for m in self.modules])

class Form(db.Model):
    __tablename__ = "form"

    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    is_filled = db.Column(db.Boolean)
    submission_date = db.Column(db.DateTime, nullable=True)
    module_id = db.Column(db.Integer, db.ForeignKey('module.id'), nullable=False) # foreign key.
    assessments = db.relationship('Assessment', backref='form', lazy=True) # one to many.
    project_id = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False) # foreign key.

    def to_dict(self):
        return dict(id=self.id,
                    is_filled=self.is_filled,
                    form_submitted_date=self.submission_date.strftime("%d/%m/%Y") if self.submission_date else "",
                    project_id=self.project_id,
                    module_id=self.module_id,
                    assessments=[a.to_dict() for a in self.assessments])

class Project(db.Model):
    __tablename__ = "project"

    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    name = db.Column(db.String(80), nullable=False)
    state = db.Column(db.String(80), nullable=False)
    create_date = db.Column(db.DateTime, nullable=False)
    due_date = db.Column(db.DateTime, nullable=False)
    sem1_bgn = db.Column(db.DateTime, nullable=False)
    sem1_end = db.Column(db.DateTime, nullable=False)
    sem2_bgn = db.Column(db.DateTime, nullable=False)
    sem2_end = db.Column(db.DateTime, nullable=False)
    exam1_bgn = db.Column(db.DateTime, nullable=False)
    exam1_end = db.Column(db.DateTime, nullable=False)
    exam2_bgn = db.Column(db.DateTime, nullable=False)
    exam2_end = db.Column(db.DateTime, nullable=False)
    forms = db.relationship('Form', backref='project', lazy=True) # one to many.

    def to_dict(self):
        semester1 = {
                "start" : self.sem1_bgn.strftime("%d/%m/%Y"),
                "end" : self.sem1_end.strftime("%d/%m/%Y"),
                "exam_period" : {
                    "start" : self.exam1_bgn.strftime("%d/%m/%Y"),
                    "end" : self.exam1_end.strftime("%d/%m/%Y")
                    }
                }
        semester2 = {
                "start" : self.sem2_bgn.strftime("%d/%m/%Y"),
                "end" : self.sem2_end.strftime("%d/%m/%Y"),
                "exam_period" : {
                    "start" : self.exam2_bgn.strftime("%d/%m/%Y"),
                    "end" : self.exam2_end.strftime("%d/%m/%Y")
                    }
                }
        return dict(project_id=self.id,
                    project_name=self.name,
                    state=str(self.state),
                    project_create=self.create_date.strftime("%d/%m/%Y"),
                    project_due=self.due_date.strftime("%d/%m/%Y"),
                    semester1=semester1,
                    semester2=semester2,
                    forms=[e.to_dict() for e in self.forms])

# ---------------- Other -------------------- #

@auth.verify_password
def verify_password(email_or_token, password):
    # first try to authenticate by token
    user = User.verify_auth_token(email_or_token)
    if not user:
        # try to authenticate with username/password
        user = User.query.filter_by(email=email_or_token).first()
        if not user or not user.verify_password(password):
            return False
    g.user = user # set global user.
    return True

# ---------------- Routes -------------------- #

@app.route('/api/hello')
@auth.login_required
def get_resource():
    return jsonify({ 'data': 'Hello, %s!' % g.user.name })

@app.route('/api/token')
@auth.login_required
def get_auth_token():
    token = g.user.generate_auth_token()
    print("Returning token", token)
    return jsonify({ 'token': token.decode('ascii') })

@app.route('/api/401')
@auth.login_required
def send_401():
    return "Not authorized", 401

@app.route('/api/module/', methods=['GET'])
@auth.login_required
def get_module():
    return jsonify({"modules":[module.to_dict() for module in Module.query.all()]})

@app.route('/api/user', methods=['GET'])
@auth.login_required
def get_user():
    return jsonify({"user":g.user.to_dict()})

@app.route('/api/project', methods=['POST'])
@auth.login_required
def create_project():
    if g.user.usertype != "ltm":
        print("Only LTM can create projects")
        return jsonify({"error_code":2,
            "error_msg":"Invalid module ID in project form"})

    # Get request params.
    json = request.get_json()
    project = Project(name=json["name"],
            state="waiting_on_academics",
            create_date=datetime.now(),
            due_date=datetime.strptime(json["due_date"], "%d/%m/%Y"),
            sem1_bgn=datetime.strptime(json["semester1"]["start"], "%d/%m/%Y"),
            sem1_end=datetime.strptime(json["semester1"]["end"], "%d/%m/%Y"),
            sem2_bgn=datetime.strptime(json["semester2"]["start"], "%d/%m/%Y"),
            sem2_end=datetime.strptime(json["semester2"]["end"], "%d/%m/%Y"),
            exam1_bgn=datetime.strptime(json["semester1"]["exam_period"]["start"], "%d/%m/%Y"),
            exam1_end=datetime.strptime(json["semester1"]["exam_period"]["end"], "%d/%m/%Y"),
            exam2_bgn=datetime.strptime(json["semester2"]["exam_period"]["start"], "%d/%m/%Y"),
            exam2_end=datetime.strptime(json["semester2"]["exam_period"]["end"], "%d/%m/%Y"))
    db.session.add(project)
    db.session.flush()
    db.session.refresh(project)
    for module_id in json["modules"]:
        form = Form(is_filled=False,
                module_id=module_id,
                project_id=project.id)
        db.session.add(form)
    db.session.commit()
    return jsonify({"error_code":0, "error_msg":"success"})

@app.route('/api/project', methods=['GET'])
@auth.login_required
def get_project():
    all_projects = Project.query.all()

    # return complete projects for LTM.
    if g.user.usertype in ("ltm", "tutor"):
        print("Returning all projects for LTM and tutor")
        return jsonify({"projects":[project.to_dict() for project in all_projects]})

    # return projects with only module forms for academic.
    elif g.user.usertype == "academic":
        print("Returning projects for user", g.user.id)

        # Find modules of current academic.
        acad_modules = [module.id for module in Module.query.filter_by(academic=g.user.id).all()]
        print("Modules of current user", acad_modules)

        ret_dict = {"projects":[project.to_dict() for project in all_projects]}

        # delete forms not relevant to current user from dict.
        for project in ret_dict["projects"]:
            new_forms = []
            for form in project["forms"]:
                if form["module_id"] in acad_modules:
                    new_forms.append(form)
            project["forms"] = new_forms

        return jsonify(ret_dict)

@app.route('/api/form', methods=['POST'])
@auth.login_required
def create_forms():
    # only academic can call this API.
    if g.user.usertype != "academic":
        return jsonify({"error_code":1,
            "error_msg":"API can only be called by academic"})

    json = request.get_json()
    if not json:
        return jsonify({"error_code":4, "error_msg":"Bad request"})

    for form_json in json["forms"]:
        result = Form.query.filter_by(id=form_json["id"]).all()
        if len(result) == 0:
            return jsonify({"error_code":3, "error_msg":"Form does not exist"})

        # gather old assessments for deletion.
        for asm in Assessment.query.filter_by(form_id=form_json["id"]).all():
            db.session.delete(asm)

        form = result[0]
        form.is_filled = True
        form.submission_date = datetime.now()
        for asm_json in form_json["assessments"]:
            Assessment(name=asm_json["asm_name"],
                    format=asm_json["asm_format"],
                    marks=asm_json["asm_per"],
                    release_date=datetime.strptime(asm_json["asm_release"], "%d/%m/%Y"),
                    submission_date=datetime.strptime(asm_json["asm_due"], "%d/%m/%Y"),
                    form=form)
        db.session.add(form)
    db.session.commit()

    # change project state if required.
    projects = Project.query.all()
    for project in projects:
        forms = Form.query.filter_by(project_id=project.id).all()
        if all([form.is_filled for form in forms]):
            if project.state == "waiting_on_academics":
                project.state = "assessment_data_collected"
                db.session.add(project)

    db.session.commit()
    return jsonify({"error_code":0, "error_msg":"success"})

@app.route('/api/assessment', methods=['PUT'])
@auth.login_required
def update_assessment():
    # only tutor can call this API.
    if g.user.usertype != "tutor":
        return jsonify({"error_code":5,
            "error_msg":"API can only be called by academic"})

    json = request.get_json()
    if not json:
        return jsonify({"error_code":4, "error_msg":"Bad request"})

    for asm_json in json["assessments"]:
        result = Assessment.query.filter_by(id=asm_json["id"]).all()
        if len(result) == 0:
            return jsonify({"error_code":3, "error_msg":"Assessment " + str(asm_json["id"]) + " does not exist"})
        asm = result[0]
        print("Updating assessment ", asm.id, ",", asm.name)
        asm.release_date = datetime.strptime(asm_json["new_start"], "%d/%m/%Y")
        asm.submission_date = datetime.strptime(asm_json["new_end"], "%d/%m/%Y")
        db.session.add(asm)

    db.session.commit()

    return jsonify({"error_code":0, "error_msg":"success"})

@app.route('/api/print/<project_id>', methods=['GET'])
@auth.login_required
def print_booklet(project_id):
    # only LTM can call this API.
    if g.user.usertype != "ltm":
        return jsonify({"error_code":5,
            "error_msg":"API can only be called by LTM"})

    result = Project.query.filter_by(id=project_id).all()
    if len(result) == 0:
        return jsonify({"error_code":3, "error_msg":"Project " + str(project_id) + " does not exist"})
    project = result[0]
    print("Printing project", project_id)

    # Fetch all modules.
    modules = Module.query.all()
    module_data = {}
    for module in modules:
        module_data[module.id] = {}
        module_data[module.id]["name"] = module.name
        module_data[module.id]["code"] = module.code
    forms = Form.query.filter_by(project_id=project_id).all()

    document = Document()
    # set page size.
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    document.add_heading("Module Assessments", 0)
    p = document.add_paragraph("Project name: " + project.name)

    for form in forms:
        document.add_heading(module_data[form.module_id]["code"] + " " + module_data[form.module_id]["name"], level=1)
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Format"
        hdr_cells[1].text = "Name"
        hdr_cells[2].text = "Marks"
        hdr_cells[3].text = "Release date"
        hdr_cells[4].text = "Deadline"
        for asm in form.assessments:
            row_cells = table.add_row().cells
            row_cells[0].text = asm.format
            row_cells[1].text = asm.name
            row_cells[2].text = str(asm.marks)
            row_cells[3].text = asm.release_date.strftime("%d/%m/%Y")
            row_cells[4].text = asm.submission_date.strftime("%d/%m/%Y")

    filename = "".join([s if s.isalnum() else "-" for s in project.name]) + ".docx"
    filepath = "/tmp/" + filename

    print("Saving", filename, "to /tmp")
    document.save(filepath)

    result = send_file(filepath,
            as_attachment=True,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            conditional=False)
    result.headers["x-suggested-filename"] = filename
    return result


app.register_blueprint(api, url_prefix="/api")

if __name__ == '__main__':
    app.run(port=8081) # run the server.
